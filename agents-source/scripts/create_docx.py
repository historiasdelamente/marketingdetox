#!/usr/bin/env python3
"""
Generador de documento Word para Voiceover Psicologico.
Combina 3 partes de voiceover (.md) en un solo documento .docx formateado.

Uso:
    python create_docx.py parte1.md parte2.md parte3.md --tema "Nombre del Tema" --output salida.docx
"""

import argparse
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def parse_metadata(content):
    """Extrae metadata del encabezado del archivo .md"""
    metadata = {}
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if line.startswith('**Tema:**'):
            metadata['tema'] = line.replace('**Tema:**', '').strip()
        elif line.startswith('**Parte:**'):
            metadata['parte'] = line.replace('**Parte:**', '').strip()
        elif line.startswith('**Caracteres:**'):
            metadata['caracteres'] = line.replace('**Caracteres:**', '').strip()
        elif line.startswith('**Palabras:**'):
            metadata['palabras'] = line.replace('**Palabras:**', '').strip()
        elif line.startswith('**Duracion estimada:**'):
            metadata['duracion'] = line.replace('**Duracion estimada:**', '').strip()
        elif line.startswith('**Hashtags sugeridos:**') or line.startswith('**Hashtags:**'):
            metadata['hashtags'] = line.split(':**')[-1].strip()
        elif line.startswith('**Titulo sugerido'):
            metadata['titulo_live'] = line.split(':**')[-1].strip()

    return metadata


def extract_content(text):
    """Extrae el contenido principal y las notas de produccion del markdown."""
    lines = text.split('\n')
    content_lines = []
    production_lines = []
    in_content = False
    in_production = False
    passed_first_separator = False

    for line in lines:
        stripped = line.strip()

        # Detectar notas de produccion
        if stripped in ('*Notas de produccion:*', 'Notas de produccion:'):
            in_production = True
            in_content = False
            continue

        if in_production:
            production_lines.append(line)
            continue

        # Detectar primer separador --- (despues de metadata)
        if stripped == '---' and not passed_first_separator:
            # Verificar si ya pasamos la metadata (campos con **)
            passed_first_separator = True
            in_content = True
            continue

        # Detectar separadores siguientes (antes de notas de produccion)
        if stripped == '---' and in_content:
            continue  # Ignorar separadores adicionales

        # Saltar metadata (titulo y campos **)
        if not in_content:
            continue

        # Agregar contenido
        content_lines.append(line)

    content = '\n'.join(content_lines).strip()
    production_notes = '\n'.join(production_lines).strip()

    # Limpiar asteriscos de markdown en contenido
    content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)  # Bold
    content = re.sub(r'\*(.+?)\*', r'\1', content)  # Italic

    return content, production_notes


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)


def create_docx(parte1_path, parte2_path, parte3_path, tema, output_path):
    """Crea el documento Word combinando las 3 partes."""

    doc = Document()

    # === CONFIGURAR MARGENES ===
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === CONFIGURAR ESTILOS ===
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5

    # === LEER LOS 3 ARCHIVOS ===
    partes_data = []
    total_chars = 0
    total_words = 0

    for path in [parte1_path, parte2_path, parte3_path]:
        with open(path, 'r', encoding='utf-8') as f:
            raw = f.read()

        metadata = parse_metadata(raw)
        content, prod_notes = extract_content(raw)

        chars = len(content)
        words = len(content.split())
        total_chars += chars
        total_words += words

        partes_data.append({
            'metadata': metadata,
            'content': content,
            'production_notes': prod_notes,
            'chars': chars,
            'words': words
        })

    # Obtener metadata general del primer archivo
    meta1 = partes_data[0]['metadata']
    titulo_live = meta1.get('titulo_live', tema)
    hashtags = meta1.get('hashtags', '#SanacionEmocional #HistoriasDeLaMente')
    duracion_total = f"{total_words // 140} minutos"

    # === PORTADA ===
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph('')

    # Titulo principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('VOICEOVER PSICOLOGICO')
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    title_run.font.name = 'Calibri'
    title_run.bold = True

    # Titulo del voiceover
    main_title_p = doc.add_paragraph()
    main_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title_run = main_title_p.add_run(titulo_live)
    main_title_run.font.size = Pt(28)
    main_title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    main_title_run.font.name = 'Calibri'
    main_title_run.bold = True

    # Linea decorativa
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run('_' * 40)
    line_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # Metadata de portada
    meta_items = [
        f'Tema: {tema}',
        f'Fecha: {datetime.now().strftime("%d de %B de %Y")}',
        f'Duracion total estimada: {duracion_total}',
        f'Caracteres totales: {total_chars:,}',
        f'Palabras totales: {total_words:,}',
    ]

    for item in meta_items:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_p.add_run(item)
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        meta_run.font.name = 'Calibri'

    # Hashtags
    doc.add_paragraph('')
    hash_p = doc.add_paragraph()
    hash_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hash_run = hash_p.add_run(hashtags)
    hash_run.font.size = Pt(10)
    hash_run.font.color.rgb = RGBColor(0x33, 0x99, 0xCC)
    hash_run.font.name = 'Calibri'
    hash_run.italic = True

    # === PARTES 1, 2 y 3 ===
    parte_nombres = [
        'PARTE 1: GANCHO + SETUP EMOCIONAL',
        'PARTE 2: CONFRONTACION + EXPLORACION PROFUNDA',
        'PARTE 3: RESOLUCION + LLAMADO A LA ACCION'
    ]

    for i, parte in enumerate(partes_data):
        # Salto de pagina
        doc.add_page_break()

        # Encabezado de parte
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_p.add_run(parte_nombres[i])
        header_run.font.size = Pt(18)
        header_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        header_run.font.name = 'Calibri'
        header_run.bold = True

        # Linea bajo el encabezado
        line_p = doc.add_paragraph()
        line_run = line_p.add_run('_' * 60)
        line_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
        line_run.font.size = Pt(8)

        # Metadata de la parte
        meta_text = f'Caracteres: {parte["chars"]:,}  |  Palabras: {parte["words"]:,}  |  Duracion: ~{parte["words"] // 140} min'
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(meta_text)
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        meta_run.font.name = 'Calibri'
        meta_run.italic = True

        doc.add_paragraph('')  # Espacio

        # Contenido del voiceover
        paragraphs = parte['content'].split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Detectar si es un titulo (empieza con #)
            if para_text.startswith('#'):
                para_text = re.sub(r'^#+\s*', '', para_text)
                p = doc.add_paragraph()
                run = p.add_run(para_text)
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                run.bold = True
                continue

            # Parrafo normal
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Manejar saltos de linea simples dentro del parrafo
            lines = para_text.split('\n')
            for j, line in enumerate(lines):
                line = line.strip()
                if line.startswith('- '):
                    line = '  ' + line
                run = p.add_run(line)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j < len(lines) - 1:
                    p.add_run('\n')

        # Notas de produccion
        if parte['production_notes']:
            doc.add_paragraph('')  # Espacio

            # Linea separadora
            sep_p = doc.add_paragraph()
            sep_run = sep_p.add_run('_' * 40)
            sep_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
            sep_run.font.size = Pt(8)

            notes_title_p = doc.add_paragraph()
            notes_title_run = notes_title_p.add_run('Notas de produccion:')
            notes_title_run.font.size = Pt(10)
            notes_title_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            notes_title_run.font.name = 'Calibri'
            notes_title_run.bold = True
            notes_title_run.italic = True

            notes_lines = parte['production_notes'].split('\n')
            for note_line in notes_lines:
                note_line = note_line.strip()
                if not note_line:
                    continue
                note_p = doc.add_paragraph()
                note_run = note_p.add_run(note_line)
                note_run.font.size = Pt(10)
                note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                note_run.font.name = 'Calibri'
                note_run.italic = True

    # === GUARDAR ===
    doc.save(output_path)
    print(f'Documento creado: {output_path}')
    print(f'Caracteres totales: {total_chars:,}')
    print(f'Palabras totales: {total_words:,}')
    print(f'Duracion estimada: {duracion_total}')


def main():
    parser = argparse.ArgumentParser(description='Genera documento Word de voiceover psicologico')
    parser.add_argument('parte1', help='Ruta al archivo .md de la Parte 1')
    parser.add_argument('parte2', help='Ruta al archivo .md de la Parte 2')
    parser.add_argument('parte3', help='Ruta al archivo .md de la Parte 3')
    parser.add_argument('--tema', required=True, help='Nombre del tema')
    parser.add_argument('--output', required=True, help='Ruta del archivo .docx de salida')

    args = parser.parse_args()

    # Verificar que los archivos existen
    for path in [args.parte1, args.parte2, args.parte3]:
        if not os.path.exists(path):
            print(f'Error: No se encontro el archivo {path}')
            sys.exit(1)

    # Crear directorio de salida si no existe
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    create_docx(args.parte1, args.parte2, args.parte3, args.tema, args.output)


if __name__ == '__main__':
    main()
